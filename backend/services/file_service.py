import os
import logging
from typing import Optional, List
from pathlib import Path

import pdfplumber
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Try to import PyMuPDF with proper error handling
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    try:
        import pymupdf as fitz
        PYMUPDF_AVAILABLE = True
    except ImportError:
        PYMUPDF_AVAILABLE = False
        logging.warning("PyMuPDF not available, using pdfplumber only")

logger = logging.getLogger(__name__)


class FileService:
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg', '.mp3', '.wav', '.m4a', '.webm', '.ogg']
        self.max_file_size = 10 * 1024 * 1024  # 10MB

    def detect_file_type(self, filepath: str) -> str:
        """Detect file type from extension. Returns 'pdf', 'docx', 'txt', 'image', or 'audio'."""
        ext = Path(filepath).suffix.lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext == '.docx':
            return 'docx'
        elif ext == '.txt':
            return 'txt'
        elif ext in ('.png', '.jpg', '.jpeg'):
            return 'image'
        elif ext in ('.mp3', '.wav', '.m4a', '.webm', '.ogg'):
            return 'audio'
        return 'unknown'

    def extract_text_universal(self, filepath: str) -> Optional[List[dict]]:
        """Route to the correct extractor based on file type. Returns [{page, text}]."""
        file_type = self.detect_file_type(filepath)
        if file_type == 'pdf':
            return self.extract_text_with_pages(filepath)
        elif file_type == 'docx':
            return self._extract_docx(filepath)
        elif file_type == 'txt':
            return self._extract_txt(filepath)
        elif file_type == 'image':
            return self._extract_image_ocr(filepath)
        else:
            logger.error(f"Unsupported file type for: {filepath}")
            return None

    def _extract_docx(self, filepath: str) -> Optional[List[dict]]:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if not paragraphs:
                return None
            # Treat each ~20 paragraphs as a "page"
            pages = []
            chunk_size = 20
            for i in range(0, len(paragraphs), chunk_size):
                text = "\n".join(paragraphs[i:i + chunk_size])
                pages.append({"page": (i // chunk_size) + 1, "text": text})
            return pages if pages else None
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return None

    def _extract_txt(self, filepath: str) -> Optional[List[dict]]:
        """Extract text from plain text file."""
        try:
            with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            if not content.strip():
                return None
            # Split into ~2000-char "pages"
            pages = []
            lines = content.split("\n")
            current_text = ""
            page_num = 1
            for line in lines:
                if len(current_text) + len(line) > 2000 and current_text:
                    pages.append({"page": page_num, "text": current_text.strip()})
                    page_num += 1
                    current_text = ""
                current_text += line + "\n"
            if current_text.strip():
                pages.append({"page": page_num, "text": current_text.strip()})
            return pages if pages else None
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            return None

    def _extract_image_ocr(self, filepath: str) -> Optional[List[dict]]:
        """Extract text from image using pytesseract OCR."""
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(filepath)
            text = pytesseract.image_to_string(img)
            if not text.strip():
                return [{"page": 1, "text": "[Image uploaded — no text detected by OCR]"}]
            return [{"page": 1, "text": text.strip()}]
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return [{"page": 1, "text": "[Image uploaded — OCR unavailable]"}]

    # ---- Existing PDF methods (kept intact) ----

    def chunking_function(self, text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        if not text:
            return []

        try:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", ". ", " ", ""],
            )
            chunks = splitter.split_text(text)
            logger.info(f"Chunked into {len(chunks)} chunks")
            return chunks

        except Exception as e:
            logger.error(f"Error chunking text: {str(e)}")
            return [text]

    def extract_text(self, filepath: str) -> Optional[str]:
        """Extract text from PDF using multiple methods for better accuracy."""
        try:
            if not self._validate_file(filepath):
                return None

            text = self._extract_with_pdfplumber(filepath)
            if not text or len(text.strip()) < 50:
                if PYMUPDF_AVAILABLE:
                    text = self._extract_with_pymupdf(filepath)

            if text:
                logger.info(f"Successfully extracted text from {filepath}")
                return text.strip()
            else:
                logger.error(f"Failed to extract text from {filepath}")
                return None

        except Exception as e:
            logger.error(f"Error extracting text from {filepath}: {str(e)}")
            return None

    def _extract_with_pdfplumber(self, filepath: str) -> Optional[str]:
        try:
            full_text = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    if text.strip():
                        full_text.append(text.strip())
            return "\n\n".join(full_text)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")
            return None

    def _extract_with_pymupdf(self, filepath: str) -> Optional[str]:
        if not PYMUPDF_AVAILABLE:
            return None
        try:
            full_text = []
            doc = fitz.open(filepath)
            for page in doc:
                text = page.get_text()
                if text.strip():
                    full_text.append(text.strip())
            doc.close()
            return "\n\n".join(full_text)
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {str(e)}")
            return None

    def extract_text_with_pages(self, filepath: str) -> Optional[list]:
        """Extract text from PDF with page numbers."""
        try:
            if not self._validate_file(filepath):
                return None

            pages = []
            try:
                with pdfplumber.open(filepath) as pdf:
                    for i, page in enumerate(pdf.pages, start=1):
                        text = page.extract_text() or ""
                        if text.strip():
                            pages.append({"page": i, "text": text.strip()})
            except Exception as e:
                logger.warning(f"pdfplumber page extraction failed: {e}")

            if not pages and PYMUPDF_AVAILABLE:
                try:
                    doc = fitz.open(filepath)
                    for i, page in enumerate(doc, start=1):
                        text = page.get_text()
                        if text.strip():
                            pages.append({"page": i, "text": text.strip()})
                    doc.close()
                except Exception as e:
                    logger.warning(f"PyMuPDF page extraction failed: {e}")

            return pages if pages else None

        except Exception as e:
            logger.error(f"Error in extract_text_with_pages: {e}")
            return None

    def chunking_function_with_pages(self, page_texts: list, chunk_size: int = 1000, chunk_overlap: int = 200) -> list:
        """Chunk text while tracking page provenance."""
        if not page_texts:
            return []

        try:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", ". ", " ", ""],
            )

            docs = [
                Document(page_content=pt["text"], metadata={"page": pt["page"]})
                for pt in page_texts
            ]
            split_docs = splitter.split_documents(docs)

            result = []
            for doc in split_docs:
                pages = sorted(set(
                    [doc.metadata["page"]] if "page" in doc.metadata else []
                ))
                result.append({"text": doc.page_content, "pages": pages})

            return result

        except Exception as e:
            logger.error(f"Error in chunking_function_with_pages: {e}")
            return [{"text": pt["text"], "pages": [pt["page"]]} for pt in page_texts]

    def _validate_file(self, filepath: str) -> bool:
        try:
            if not os.path.exists(filepath):
                logger.error(f"File not found: {filepath}")
                return False

            file_size = os.path.getsize(filepath)
            if file_size > self.max_file_size:
                logger.error(f"File too large: {file_size} bytes")
                return False

            file_ext = Path(filepath).suffix.lower()
            if file_ext not in self.supported_extensions:
                logger.error(f"Unsupported file type: {file_ext}")
                return False

            return True

        except Exception as e:
            logger.error(f"File validation error: {str(e)}")
            return False

    def get_file_info(self, filepath: str) -> Optional[dict]:
        """Get basic information about a file."""
        try:
            file_size = os.path.getsize(filepath)
            file_name = os.path.basename(filepath)
            file_type = self.detect_file_type(filepath)

            page_count = 0
            if file_type == 'pdf' and PYMUPDF_AVAILABLE:
                try:
                    doc = fitz.open(filepath)
                    page_count = len(doc)
                    doc.close()
                except Exception:
                    page_count = 0

            return {
                "filename": file_name,
                "size_bytes": file_size,
                "size_mb": round(file_size / (1024 * 1024), 2),
                "page_count": page_count,
                "file_type": file_type,
            }

        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return None
